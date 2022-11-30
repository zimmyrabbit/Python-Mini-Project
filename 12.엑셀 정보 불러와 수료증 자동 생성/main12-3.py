import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document(r'12.엑셀 정보 불러와 수료증 자동 생성\수료증양식.docx')

style = doc.styles['Normal']
style.font.name = '맑은고딕'
style._element.rPr.rFonts.set(qn('w:esatAsia'), '맑은고딕')
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('            제 2022-0001 호\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('수 료 증')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('      성 명 : 이진채\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)

run = para.add_run('      생 년 월 일 : 1997.01.01\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)

run = para.add_run('      교 육 과 정 : 파이썬\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)

run = para.add_run('      교 육 날 짜 : 2022.11.21 ~ 2022.11.30\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('      파이썬으로 수료증 만들기\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)

run = para.add_run('      테스트 입니다.\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('2022.11.30\n')
run.font.name = '맑은고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('이진채\n')
run.font.name = '맑은고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'12.엑셀 정보 불러와 수료증 자동 생성\수료증결과.docx')