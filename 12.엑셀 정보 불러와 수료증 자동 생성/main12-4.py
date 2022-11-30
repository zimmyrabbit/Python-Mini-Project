from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

load_wb = load_workbook(r'12.엑셀 정보 불러와 수료증 자동 생성\수료증명단.xlsx')
load_ws = load_wb.active

name_list = []
birthday_list = []
ho_list = []

for i in range(1, load_ws.max_row + 1):
    name_list.append(load_ws.cell(i,1).value)
    birthday_list.append(load_ws.cell(i,2).value)
    ho_list.append(load_ws.cell(i,3).value)

for i in range(len(name_list)):
    doc = docx.Document(r'12.엑셀 정보 불러와 수료증 자동 생성\수료증양식.docx')

    style = doc.styles['Normal']
    style.font.name = '맑은고딕'
    style._element.rPr.rFonts.set(qn('w:esatAsia'), '맑은고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph() 
    run = para.add_run('\n\n')
    run = para.add_run('            제 ' + ho_list[i] + ' 호\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph() 
    run = para.add_run('\n\n')
    run = para.add_run('수 료 증')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph() 
    run = para.add_run('\n\n')
    run = para.add_run('      성 명 : ' + name_list[i] + '\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run('      생 년 월 일 : ' + birthday_list[i] + '\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run('      교 육 과 정 : 파이썬\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run('      교 육 날 짜 : 2022.11.21 ~ 2022.11.30\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph() 
    run = para.add_run('\n\n')
    run = para.add_run('      파이썬으로 수료증 만들기\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run('      테스트 입니다.\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph() 
    run = para.add_run('\n\n')
    run = para.add_run('2022.11.30\n')
    run.font.name = '맑은고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph() 
    run = para.add_run('\n\n')
    run = para.add_run('이진채\n')
    run.font.name = '맑은고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(r'12.엑셀 정보 불러와 수료증 자동 생성\\' + name_list[i] + '.docx')
    convert('12.엑셀 정보 불러와 수료증 자동 생성\\' + name_list[i] + '.docx', '12.엑셀 정보 불러와 수료증 자동 생성\\' + name_list[i] + '.pdf')